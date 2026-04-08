"""
Interview Agent Universal — app.py
Generalized real-time AI interview co-pilot.
Profile, briefings, and API key are all stored locally — nothing is hardcoded.
"""

import os
import io
import json
import time
import ipaddress
from collections import defaultdict
from pathlib import Path
from urllib.parse import urlparse

import anthropic
import requests as req_lib
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, Response, stream_with_context
from werkzeug.utils import secure_filename

load_dotenv()

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024  # 20 MB upload limit

PROFILE_PATH = Path("profile.json")
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# ---------------------------------------------------------------------------
# CLIENT (initialized lazily so the app starts even without a key)
# ---------------------------------------------------------------------------
_client = None

def get_client():
    global _client
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if key and (_client is None):
        _client = anthropic.Anthropic(api_key=key)
    return _client


# ---------------------------------------------------------------------------
# RATE LIMITING (in-memory, per IP)
# ---------------------------------------------------------------------------
_rate_store: dict = defaultdict(list)

def _check_rate(key: str, max_calls: int = 30, window: int = 60) -> bool:
    now = time.time()
    calls = _rate_store[key]
    calls[:] = [t for t in calls if now - t < window]
    if len(calls) >= max_calls:
        return False
    calls.append(now)
    return True

def rate_limit(max_calls: int = 30, window: int = 60):
    """Decorator — rate-limit by IP."""
    def decorator(f):
        def wrapper(*args, **kwargs):
            ip = request.remote_addr or "unknown"
            if not _check_rate(f"{f.__name__}:{ip}", max_calls, window):
                return jsonify({"error": "Rate limit exceeded. Please slow down."}), 429
            return f(*args, **kwargs)
        wrapper.__name__ = f.__name__
        return wrapper
    return decorator


# ---------------------------------------------------------------------------
# SECURITY HEADERS
# ---------------------------------------------------------------------------
@app.after_request
def security_headers(resp):
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["X-Frame-Options"] = "SAMEORIGIN"
    resp.headers["X-XSS-Protection"] = "1; mode=block"
    resp.headers["Referrer-Policy"] = "no-referrer"
    resp.headers["Permissions-Policy"] = "geolocation=(), camera=()"
    return resp


# ---------------------------------------------------------------------------
# PROFILE HELPERS
# ---------------------------------------------------------------------------
def load_profile() -> dict:
    if PROFILE_PATH.exists():
        try:
            return json.loads(PROFILE_PATH.read_text())
        except Exception:
            pass
    return {}

def save_profile(data: dict):
    PROFILE_PATH.write_text(json.dumps(data, indent=2))

def profile_to_text(profile: dict) -> str:
    if not profile:
        return ""
    lines = []
    if profile.get("name"):
        lines.append(f"Name: {profile['name']}")
    if profile.get("title"):
        lines.append(f"Title: {profile['title']}")
    if profile.get("organization"):
        lines.append(f"Organization: {profile['organization']}")
    if profile.get("location"):
        lines.append(f"Location: {profile['location']}")
    if profile.get("summary"):
        lines.append(f"\nSummary: {profile['summary']}")
    if profile.get("expertise"):
        lines.append(f"\nExpertise: {', '.join(profile['expertise'])}")
    if profile.get("experience"):
        lines.append("\nExperience:")
        for exp in profile["experience"][:6]:
            lines.append(f"  - {exp.get('title','')} at {exp.get('org','')} ({exp.get('dates','')})")
            for h in exp.get("highlights", [])[:2]:
                lines.append(f"      • {h}")
    if profile.get("education"):
        lines.append("\nEducation:")
        for ed in profile["education"]:
            lines.append(f"  - {ed.get('degree','')} — {ed.get('institution','')} ({ed.get('year','')})")
    if profile.get("publications"):
        lines.append("\nPublications/Writing (selected):")
        for pub in profile["publications"][:5]:
            lines.append(f"  - {pub}")
    if profile.get("achievements"):
        lines.append("\nKey achievements:")
        for a in profile["achievements"][:4]:
            lines.append(f"  - {a}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# SYSTEM PROMPT (built dynamically from stored profile)
# ---------------------------------------------------------------------------
BASE_SYSTEM_PROMPT = """You are a real-time interview assistant for {name}. You are a paid expert giving authoritative, specific, data-rich answers on any topic relevant to their domain — drawing on their credentials and your own deep expertise.

## WHO YOU ARE ASSISTING:
{profile_text}

## EXPERTISE AREAS (from profile + general knowledge):
{expertise_block}

## WHEN A SESSION BRIEFING IS PROVIDED:
Prioritize it as your primary reference. Cite its specific numbers, provisions, and details directly.

## RESPONSE FORMAT (STRICT):
- 2 to 4 bullet points MAXIMUM
- **Bold** the KEY TERM or DATA POINT at the start of each bullet
- Use specific numbers, legislation names, dollar figures, dates — from the briefing if available, from your expertise otherwise
- No hyperlinks in live bullets (those go in the post-call summary only)
- Always give a substantive response — even if the transcript is partial, use whatever context exists

## TONE AND CONTENT:
- Drive the conversation with facts, legal arguments, technical specifics, and policy narratives
- Lead with the substance — statutes, numbers, court doctrine, regulatory outcomes, real-world examples
- Do NOT reference the person's past work history or proceedings unless directly asked about their background
- State facts directly — not "they testified that" or "their work showed" — just the facts
- You are a paid expert giving opinion and analysis. Be direct, fluid, and authoritative

## NON-NEGOTIABLE:
- NEVER say "I don't have information on" or any deflecting hedge
- NEVER deflect — always give specific, substantive content
- SHORT — read at a glance mid-conversation
"""

def build_system_prompt(profile: dict) -> str:
    name = profile.get("name", "the user") if profile else "the user"
    profile_text = profile_to_text(profile) if profile else "No profile configured. Provide general expert assistance."
    expertise = profile.get("expertise", [])
    expertise_block = (
        "\n".join(f"- {e}" for e in expertise)
        if expertise
        else "- General professional expertise (build your profile to customize this)"
    )
    return BASE_SYSTEM_PROMPT.format(
        name=name,
        profile_text=profile_text,
        expertise_block=expertise_block,
    )


# ---------------------------------------------------------------------------
# SSRF PROTECTION
# ---------------------------------------------------------------------------
def is_safe_url(url: str) -> bool:
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            return False
        host = parsed.hostname or ""
        if not host:
            return False
        blocked = {"localhost", "0.0.0.0", "::1"}
        if host.lower() in blocked:
            return False
        try:
            ip = ipaddress.ip_address(host)
            if any([ip.is_private, ip.is_loopback, ip.is_reserved,
                    ip.is_link_local, ip.is_multicast]):
                return False
        except ValueError:
            pass  # hostname, not raw IP — OK
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# ROUTES — SETUP
# ---------------------------------------------------------------------------
@app.route("/setup")
def setup():
    return render_template("setup.html")


@app.route("/check-setup")
def check_setup():
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    profile = load_profile()
    return jsonify({
        "has_key": bool(key and len(key) > 20),
        "has_profile": bool(profile.get("name")),
    })


@app.route("/save-key", methods=["POST"])
@rate_limit(max_calls=10, window=60)
def save_key():
    data = request.get_json(silent=True) or {}
    key = data.get("key", "").strip()
    if not key.startswith("sk-ant-"):
        return jsonify({"error": "Invalid API key format. Expected sk-ant-..."}), 400
    if len(key) < 40:
        return jsonify({"error": "Key too short — please check and try again."}), 400
    # Verify key against Claude
    try:
        test = anthropic.Anthropic(api_key=key)
        test.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=5,
            messages=[{"role": "user", "content": "hi"}],
        )
    except anthropic.AuthenticationError:
        return jsonify({"error": "Authentication failed. Check your API key."}), 401
    except Exception as e:
        return jsonify({"error": f"Could not verify key: {str(e)[:100]}"}), 400
    # Save to .env
    env_path = Path(".env")
    env_path.write_text(f"ANTHROPIC_API_KEY={key}\n")
    os.environ["ANTHROPIC_API_KEY"] = key
    global _client
    _client = anthropic.Anthropic(api_key=key)
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# ROUTES — PROFILE
# ---------------------------------------------------------------------------
@app.route("/profile", methods=["GET"])
def get_profile():
    return jsonify(load_profile())


@app.route("/profile", methods=["POST"])
@rate_limit(max_calls=20, window=60)
def post_profile():
    data = request.get_json(silent=True) or {}
    save_profile(data)
    return jsonify({"ok": True})


@app.route("/build-profile", methods=["POST"])
@rate_limit(max_calls=5, window=60)
def build_profile():
    client = get_client()
    if not client:
        return jsonify({"error": "API key not configured. Visit /setup first."}), 401

    combined_text = ""
    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "No files uploaded."}), 400

    for f in files[:10]:  # max 10 files
        fname = secure_filename(f.filename)
        ext = Path(fname).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            continue
        raw = f.read()
        if ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(raw)) as pdf:
                    combined_text += "\n\n".join(p.extract_text() or "" for p in pdf.pages[:30])
            except Exception:
                pass
        elif ext == ".docx":
            try:
                import docx as docx_lib
                doc = docx_lib.Document(io.BytesIO(raw))
                combined_text += "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception:
                pass
        elif ext == ".txt":
            combined_text += raw.decode("utf-8", errors="ignore")

    if not combined_text.strip():
        return jsonify({"error": "Could not extract text from uploaded files."}), 400

    prompt = f"""Parse these professional documents and return a JSON profile. Return ONLY valid JSON, no markdown, no code blocks.

Required fields:
{{
  "name": "Full name",
  "title": "Current job title",
  "organization": "Current employer or affiliation",
  "location": "City, State/Country",
  "summary": "3-4 sentence professional summary in first person",
  "expertise": ["area1", "area2", "area3"],
  "experience": [
    {{"title": "...", "org": "...", "dates": "...", "highlights": ["...", "..."]}}
  ],
  "education": [
    {{"degree": "...", "institution": "...", "year": "..."}}
  ],
  "publications": ["citation or title 1", "citation or title 2"],
  "skills": ["skill1", "skill2"],
  "achievements": ["achievement1", "achievement2"]
}}

DOCUMENTS:
{combined_text[:15000]}"""

    try:
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw_json = resp.content[0].text.strip()
        # Strip markdown code blocks if present
        if raw_json.startswith("```"):
            raw_json = raw_json.split("```")[1]
            if raw_json.startswith("json"):
                raw_json = raw_json[4:]
        profile = json.loads(raw_json)
        save_profile(profile)
        return jsonify({"profile": profile})
    except json.JSONDecodeError:
        return jsonify({"error": "Could not parse profile from documents. Try uploading clearer files."}), 400
    except Exception as e:
        return jsonify({"error": str(e)[:150]}), 500


# ---------------------------------------------------------------------------
# ROUTES — MAIN
# ---------------------------------------------------------------------------
@app.route("/")
def index():
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not key or len(key) < 20:
        from flask import redirect
        return redirect("/setup")
    return render_template("index.html")


@app.route("/assist-stream", methods=["POST"])
@rate_limit(max_calls=60, window=60)
def assist_stream():
    client = get_client()
    if not client:
        return jsonify({"error": "API key not configured."}), 401

    data = request.get_json(silent=True) or {}
    transcript = data.get("transcript", "").strip()
    session_briefing = data.get("briefing", "").strip()
    bad_points = [str(p)[:200] for p in data.get("badPoints", [])[:15]]
    good_points = [str(p)[:200] for p in data.get("goodPoints", [])[:8]]

    if not transcript:
        return Response(stream_with_context(iter([])), mimetype="text/event-stream")

    profile = load_profile()
    system = build_system_prompt(profile)

    context_block = f"SESSION BRIEFING:\n{session_briefing}\n\n---\n" if session_briefing else ""
    user_msg = f"{context_block}Conversation excerpt:\n\n{transcript}\n\nProvide concise talking points."

    if bad_points:
        user_msg += "\n\nDO NOT repeat or rephrase these (marked not useful):\n" + "\n".join(f"- {p}" for p in bad_points)
    if good_points:
        user_msg += "\n\nUser found these angles useful — similar style welcome:\n" + "\n".join(f"- {p}" for p in good_points)

    def generate():
        try:
            with client.messages.stream(
                model="claude-sonnet-4-6",
                max_tokens=600,
                system=system,
                messages=[{"role": "user", "content": user_msg}],
            ) as stream:
                for text in stream.text_stream:
                    yield f"data: {json.dumps(text)}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps(f'Error: {str(e)[:100]}')}\n\n"
            yield "data: [DONE]\n\n"

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/summary", methods=["POST"])
@rate_limit(max_calls=10, window=60)
def summary():
    client = get_client()
    if not client:
        return jsonify({"error": "API key not configured."}), 401

    data = request.get_json(silent=True) or {}
    full_transcript = data.get("transcript", "").strip()
    session_briefing = data.get("briefing", "").strip()
    session_name = str(data.get("session_name", "Session"))[:100]

    if not full_transcript:
        return jsonify({"summary": "No transcript recorded."})

    profile = load_profile()
    system = build_system_prompt(profile)

    brief_block = f"SESSION BRIEFING:\n{session_briefing}\n\n---\n" if session_briefing else ""
    prompt = f"""Session: {session_name}

{brief_block}FULL TRANSCRIPT:
{full_transcript}

Provide a structured post-session summary with these sections:
1. KEY TOPICS COVERED
2. STRONGEST MOMENTS — where the speaker demonstrated expertise or cited strong data
3. GAPS OR MISSED OPPORTUNITIES
4. FOLLOW-UP ACTIONS
5. ALL REFERENCED SOURCES WITH HYPERLINKS

Format in clean markdown."""

    try:
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2000,
            system=system,
            messages=[{"role": "user", "content": prompt}],
        )
        return jsonify({"summary": resp.content[0].text})
    except Exception as e:
        return jsonify({"summary": f"Error: {str(e)[:150]}"}), 500


@app.route("/extract", methods=["POST"])
@rate_limit(max_calls=20, window=60)
def extract():
    text = ""

    if request.is_json:
        url = (request.get_json(silent=True) or {}).get("url", "").strip()
        if not url:
            return jsonify({"error": "No URL provided."}), 400
        if not is_safe_url(url):
            return jsonify({"error": "URL not allowed (private/local addresses are blocked)."}), 400
        try:
            from bs4 import BeautifulSoup
            r = req_lib.get(
                url, timeout=12,
                headers={"User-Agent": "Mozilla/5.0 (compatible; InterviewAgent/1.0)"},
                allow_redirects=True,
            )
            r.raise_for_status()
            soup = BeautifulSoup(r.text, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe"]):
                tag.decompose()
            raw = soup.get_text(separator="\n", strip=True)
            lines = [l.strip() for l in raw.splitlines() if len(l.strip()) > 40]
            text = "\n".join(lines)[:12000]
            text = f"[Source: {url}]\n\n{text}"
        except Exception as e:
            return jsonify({"error": f"Could not fetch URL: {str(e)[:120]}"}), 400

    elif "file" in request.files:
        f = request.files["file"]
        fname = secure_filename(f.filename)
        ext = Path(fname).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            return jsonify({"error": f"File type '{ext}' not supported. Use PDF, DOCX, or TXT."}), 400
        raw = f.read()

        if ext == ".pdf":
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(raw)) as pdf:
                    pages = [p.extract_text() or "" for p in pdf.pages[:40]]
                text = "\n\n".join(p for p in pages if p.strip())
                text = f"[Source: {fname}]\n\n{text}"
            except Exception as e:
                return jsonify({"error": f"PDF extraction failed: {str(e)[:100]}"}), 400

        elif ext == ".docx":
            try:
                import docx as docx_lib
                doc = docx_lib.Document(io.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                text = f"[Source: {fname}]\n\n{text}"
            except Exception as e:
                return jsonify({"error": f"DOCX extraction failed: {str(e)[:100]}"}), 400

        elif ext == ".txt":
            text = raw.decode("utf-8", errors="ignore")
            text = f"[Source: {fname}]\n\n{text}"
    else:
        return jsonify({"error": "No file or URL provided."}), 400

    if not text.strip():
        return jsonify({"error": "No text could be extracted from this source."}), 400

    return jsonify({"text": text.strip()})


if __name__ == "__main__":
    app.run(debug=False, port=5055, host="127.0.0.1")
